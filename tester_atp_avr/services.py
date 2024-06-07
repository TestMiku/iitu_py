import abc
import collections.abc
import dataclasses
import datetime
import io
import logging
import os
import pathlib
import re
import shutil
import tempfile
import traceback
import typing
import uuid
from dataclasses import asdict, dataclass
from functools import cached_property
from types import SimpleNamespace
from typing import Tuple, List, Any, Dict

import bs4
import camelot
import docx.document
import docx.enum.style
import docx.enum.table
import openpyxl
import openpyxl.cell.cell
import openpyxl.worksheet.worksheet
import pandas as pd
import PyPDF2
from camelot.core import TableList
from django.core.files.storage import Storage, default_storage
from django.db import models
from django.template import Template
from django.template.loader import get_template
from django.utils.translation import gettext_lazy as _
from docx import Document
from Levenshtein import distance
from openpyxl.reader.excel import load_workbook
from openpyxl.styles import Alignment, Font
from openpyxl.styles.borders import BORDER_THIN, Border, Side
from pprint import pprint

from avh_services.formatters import (
    format_count,
    format_number,
    format_number_as_words,
    remove_spaces,
)
from avh_services.keys import tcp_key
from tester_atp_avr.models import TCP, TCPCategory, TCPFile

from .models import ContratcNumberAndDate, WorkType


class ServiceError(Exception):
    pass


class TableNotFoundError(ServiceError):
    pass


class AVRTableNotFoundError(ServiceError):
    pass


class ATPTableNotFoundError(ServiceError):
    pass


@dataclass(kw_only=True)
class ComparerAndProcessor:
    avr_rows: dict[str, "AVRRow"]
    atp_rows: dict[str, "ATPRow"]
    difference_sum: int | None = None
    output_path: pathlib.Path
    vat: int = 12
    storage: Storage = default_storage
    font_name: str = "Times New Roman"
    font_size: int = 12

    bs_order: str | None = None
    bs_date: str | None = None
    bs_number: str | None = None
    bs_name: str | None = None
    bs1_name: str | None = None
    akt_montaz: list = None
    akt_montaz1: list = None
    akt_montaz2: list = None
    smeta: io.BytesIO = None

    contract_number_and_date: str | None = None
    work_type: str | None = None

    @cached_property
    def _output_absolute_path(self) -> pathlib.Path:
        return pathlib.Path(self.storage.path(self.output_path))

    def compare_and_process(self) -> dict[str, str]:
        changes = []
        first_sum = sum(avr.price for avr in self.avr_rows.values())

        for avr_tcp, avr in self.avr_rows.items():
            if avr_tcp in self.atp_rows:
                atp = self.atp_rows[avr_tcp]
                if avr.count != atp.count:
                    current_price = float(avr.price)
                    changed_price = current_price / float(avr.count) * float(atp.count)
                    if current_price > changed_price:
                        changes.append(
                            f"Пункт {avr_tcp} ({avr.name}) - расход уменьшен"
                        )
                    else:
                        changes.append(
                            f"Пункт {avr_tcp} ({avr.name}) - расход увеличен"
                        )
                    avr.price = changed_price
                    avr.count = format_count(atp.count) + "*"
            else:
                if avr.added:
                    changes.append(f"Пункт {avr_tcp} ({avr.name}) - добавлен")
                else:
                    avr.price = 0
                    avr.count = format_count(avr.count) + "*"
                    changes.append(f"Пункт {avr_tcp} ({avr.name}) - исключен")
                    avr.classes = "table-danger"

        last_sum = sum(avr.price for avr in self.avr_rows.values())

        for atp_tcp, atp in self.atp_rows.items():
            if atp_tcp not in self.avr_rows:
                changes.append(f"Пункт {atp_tcp} ({atp.name}) - добавлен")
                self.avr_rows[atp_tcp] = AVRRow(**asdict(atp), price=0, added=False)
        total = last_sum
        total_vat = total * (self.vat / 100)
        total_with_vat = total + total_vat
        if self.difference_sum is None:
            order_difference = (first_sum - last_sum) * (1 + (self.vat / 100))
        else:
            order_difference = self.difference_sum - total_with_vat
        shutil.rmtree(self.storage.path(self.output_path), ignore_errors=True)
        self._output_absolute_path.mkdir(parents=True)

        data: ConvertData = dict(
            order_difference=format_number(order_difference),
            total=format_number(total),
            total_vat=format_number(total_vat),
            total_with_vat=format_number(total + total_vat),
            vat=self.vat,
            changes=changes,
            font_name=self.font_name,
            font_size=self.font_size,
            avr_rows=self.avr_rows,
            rows=[
                     [
                         "№ п/п",
                         "№ п. ТЦП",
                         "НАИМЕНОВАНИЕ И ПОСЛЕДОВАТЕЛЬНОСТЬ ВЫПОЛНЕНИЯ ЭТАПОВ РАБОТ (ВИДОВ РАБОТ)",
                         "Ед. изм",
                         "Кол-во",
                         "Полная стоимость работ в тенге, без НДС",
                         "Сроки начала и окончания работ",
                     ]
                 ]
                 + [
                     [
                         str(index),
                         avr.tcp,
                         avr.name,
                         avr.measuring_unit,
                         avr.count
                         if isinstance(avr.count, str)
                         else format_count(avr.count),
                         format_number(avr.price),
                         avr.due_date
                         if isinstance(avr.due_date, str)
                         else avr.due_date.strftime("%d.%m.%Y"),
                     ]
                     for index, avr in enumerate(
                    sorted(
                        self.avr_rows.values(),
                        key=lambda x: tcp_key(x.tcp),
                    ),
                    1,
                )
                 ],
            totals_as_words=[
                f"Всего общая стоимость работ: {format_number_as_words(total + total_vat)}, включая НДС {self.vat}%",
            ],
            bs_order=self.bs_order,
            bs_date=self.bs_date,
            bs_number=self.bs_number,
            bs_name=self.bs_name,
            bs1_name=self.bs1_name,
            akt_montaz=self.akt_montaz,
            akt_montaz1=self.akt_montaz1,
            akt_montaz2=self.akt_montaz2,
            smeta=self.smeta,
            contract_number_and_date=self.contract_number_and_date,
            work_type=self.work_type,
        )

        excel_name = str(uuid.uuid4()) + ".xlsx"
        ExcelConverter().convert(self._output_absolute_path / excel_name, data)
        word_name = str(uuid.uuid4()) + ".docx"
        WordConverter().convert(self._output_absolute_path / word_name, data)
        html_name = str(uuid.uuid4()) + ".html"
        HTMLConverter().convert(self._output_absolute_path / html_name, data)
        word_format1_name = str(uuid.uuid4()) + "_format1.docx"
        WordFormat1Converter().convert(
            self._output_absolute_path / word_format1_name, data
        )
        word_format2_name = str(uuid.uuid4()) + "_format2.docx"
        WordFormat2Converter().convert(
            self._output_absolute_path / word_format2_name, data
        )
        # docx_format1_name = str(uuid.uuid4()) + "_other_format.docx"
        # DOCXOtherFormat().convert(
        #     self._output_absolute_path / docx_format1_name, data
        # )

        return {
            "html": self.storage.url(self.output_path / html_name),
            "xlsx": self.storage.url(self.output_path / excel_name),
            "docx": self.storage.url(self.output_path / word_name),
            "docx_format1": self.storage.url(self.output_path / word_format1_name),
            "docx_format2": self.storage.url(self.output_path / word_format2_name)
            # "docx_format1": self.storage.url(
            #     self.output_path / word_name
            # ),
        }


@dataclass(kw_only=True)
class WordTemplateProcessor:
    output_path: pathlib.Path
    storage: Storage = default_storage

    @cached_property
    def _output_absolute_path(self) -> pathlib.Path:
        return pathlib.Path(self.storage.path(self.output_path))

    def process_template(self, context: dict) -> dict[str, str]:
        from docxtpl import DocxTemplate
        from django.templatetags.static import static
        from django.conf import settings
        context["current_date"]= get_current_date()
        pprint(context)
        # Path to the DOCX template
        template_path = settings.BASE_DIR / static("docx_templates/format_work_and_materials.docx").lstrip('/')
        if not template_path.exists():
            raise FileNotFoundError("The template file format_work_and_materials.docx was not found.")

        # Render the DOCX template
        doc = DocxTemplate(template_path)
        doc.render(context)

        # Generate a unique directory for the output files
        unique_id = str(uuid.uuid4())
        output_dir = os.path.join('p1', unique_id)
        absolute_output_dir = os.path.join(settings.MEDIA_ROOT, output_dir)
        os.makedirs(absolute_output_dir, exist_ok=True)

        # Save the rendered DOCX file
        word_name = f"{unique_id}.docx"
        word_output_path = os.path.join(output_dir, word_name)
        absolute_word_output_path = os.path.join(absolute_output_dir, word_name)
        doc.save(absolute_word_output_path)

        # Path to the HTML template
        html_template_path = settings.BASE_DIR / 'templates/p1/tester_atp_avr/electro_montazh.html'
        if not html_template_path.exists():
            raise FileNotFoundError("The HTML template file electro_montazh.html was not found.")

        # Render the HTML template
        html_template = get_template('p1/tester_atp_avr/electro_montazh.html')
        html_content = html_template.render(context)

        # Save the rendered HTML content to a file
        html_name = f"{unique_id}.html"
        html_output_path = os.path.join(output_dir, html_name)
        absolute_html_output_path = os.path.join(absolute_output_dir, html_name)
        with open(absolute_html_output_path, 'w', encoding='utf-8') as html_file:
            html_file.write(html_content)

        return {
            "html": self.storage.url(html_output_path),
            "electrical_installation": self.storage.url(word_output_path)
        }


class ConvertData(typing.TypedDict):
    rows: list[list[str]]
    vat: int
    total: str
    total_vat: str
    total_with_vat: str
    totals_as_words: list[str]
    changes: list[str]
    font_name: str
    font_size: int
    order_difference: str
    avr_rows: dict[str, "AVRRow"]


class Converter(abc.ABC):
    @abc.abstractmethod
    def convert(self, path: pathlib.Path, /, data: ConvertData) -> None:
        pass


class ExcelConverter(Converter):
    @cached_property
    def _workbook(self) -> openpyxl.Workbook:
        return openpyxl.Workbook()

    @cached_property
    def _worksheet(self) -> openpyxl.worksheet.worksheet.Worksheet:
        return self._workbook.active

    def _cell(
            self,
            row: int,
            column: int,
            value: typing.Any,
            is_header: bool = False,
            horizontal_center: bool = False,
            border: bool = False,
            wrap_text: bool = True,
            font_size: int | None = None,
            font_name: str | None = None,
            with_number_format: bool = False,
    ) -> openpyxl.cell.Cell:
        cell = self._worksheet.cell(row=row, column=column, value=value)
        cell.alignment = Alignment(
            vertical="center",
            horizontal="center" if horizontal_center or is_header else None,
            wrap_text=wrap_text,
        )
        cell.font = Font(
            name=font_name,
            size=font_size,
            bold=is_header,
        )
        if border:
            cell.border = Border(
                left=Side(BORDER_THIN),
                right=Side(BORDER_THIN),
                top=Side(BORDER_THIN),
                bottom=Side(BORDER_THIN),
            )
        if with_number_format:
            cell.number_format = "#,##0.00"
        return cell

    def convert(self, path: pathlib.Path, /, data: ConvertData) -> openpyxl.Workbook:
        self._cell(1, column=1, value="")
        for i, row in enumerate(data["rows"], 1):
            for j, value in enumerate(row, 1):
                self._cell(
                    row=i,
                    column=j,
                    value=value,
                    is_header=i == 1,
                    horizontal_center=j != 3,
                    border=True,
                    font_size=14 if i == 1 else data["font_size"],
                    with_number_format=i != 1 and j == 6,
                )

        column_widths = [10, 10, 50, 10, 10, 15, 10]
        for index, width in enumerate(column_widths, start=1):
            self._worksheet.column_dimensions[
                openpyxl.utils.get_column_letter(index)
            ].width = width

        max_row = self._worksheet.max_row

        self._cell(row=max_row + 2, column=5, value="Итого:")
        self._cell(
            row=max_row + 2,
            column=6,
            value=data["total"],
            horizontal_center=True,
            with_number_format=True,
        )
        self._cell(row=max_row + 3, column=5, value=f"НДС {data['vat']}%:")
        self._cell(
            row=max_row + 3,
            column=6,
            value=data["total_vat"],
            horizontal_center=True,
            with_number_format=True,
        )
        self._cell(row=max_row + 4, column=5, value="Всего с учетом НДС:")
        self._cell(
            row=max_row + 4,
            column=6,
            value=data["total_with_vat"],
            horizontal_center=True,
            with_number_format=True,
        )
        self._cell(row=max_row + 5, column=5, value="Разница с заказом:")
        self._cell(
            row=max_row + 5,
            column=6,
            value=data["order_difference"],
            horizontal_center=True,
            with_number_format=True,
        )
        for index, total_as_words in enumerate(data["totals_as_words"]):
            self._cell(
                row=max_row + 7 + index, column=5, value=total_as_words, wrap_text=False
            )
        for index, change in enumerate(data["changes"]):
            self._cell(
                row=max_row + 11 + index, column=5, value=change, wrap_text=False
            )
        self._workbook.save(path)
        return self._workbook


class WordConverter(Converter):
    def convert(self, path: pathlib.Path, /, data: ConvertData) -> None:
        document: docx.document.Document = Document()
        rows = data["rows"]
        table = document.add_table(rows=len(rows), cols=7)
        table.style = "Table Grid"
        for i, row in enumerate(rows):
            for j, column in enumerate(row):
                table.rows[i].cells[j].text = str(column)
        document.add_paragraph()
        table = document.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
        table.rows[0].cells[0].text = "Итого:"
        table.rows[0].cells[1].text = data["total"]
        table.rows[1].cells[0].text = f"НДС {data['vat']}%"
        table.rows[1].cells[1].text = data["total_vat"]
        table.rows[2].cells[0].text = "Всего с учетом НДС:"
        table.rows[2].cells[1].text = data["total_with_vat"]
        table.rows[3].cells[0].text = "Разница с заказом:"
        table.rows[3].cells[1].text = data["order_difference"]
        document.add_paragraph()
        for total_as_words in data["totals_as_words"]:
            document.add_paragraph(total_as_words)
        document.add_paragraph()
        for change in data["changes"]:
            document.add_paragraph(change)
        document.save(path)


class WordFormat1Converter(Converter):
    def convert(self, path: pathlib.Path, /, data: ConvertData) -> None:
        try:
            data["rows"] = data["rows"][1:]
            data["current_date"] = get_current_date()
            data["changes_in_smeta"] = False
            variants = ("сметному", "19.1", "19.2", "19.3", "20.1")
            if data["changes"]:
                for change in data["changes"]:
                    for variant in variants:
                        if variant in f'{change}':
                            data["changes_in_smeta"] = True

            from django.conf import settings
            from django.templatetags.static import static
            from docxtpl import DocxTemplate

            start_file_path = "/".join(
                os.path.abspath(__file__).replace("\\", "/").split("/")[:-2]
            ) + static("docx_templates/format1_start.docx")

            end_file_path = "/".join(
                os.path.abspath(__file__).replace("\\", "/").split("/")[:-2]
            ) + static("docx_templates/format1_end.docx")

            doc = DocxTemplate(start_file_path)
            data["br"] = "\n"
            doc.render(data)
            doc.save(f"{str(path)[:-5]}start.docx")

            doc = DocxTemplate(end_file_path)
            doc.render(data)
            doc.save(f"{str(path)[:-5]}end.docx")

            combine_docx(
                f"{str(path)[:-5]}start.docx",
                f"{str(path)[:-5]}end.docx",
                data["smeta"],
                path,
            )

        except:
            traceback.print_exc()
            document: docx.document.Document = Document()
            rows = data["rows"]
            table = document.add_table(rows=len(rows), cols=7)
            table.style = "Table Grid"
            for i, row in enumerate(rows):
                for j, column in enumerate(row):
                    table.rows[i].cells[j].text = str(column)
            document.add_paragraph()
            table = document.add_table(rows=4, cols=2)
            table.style = "Table Grid"
            table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
            table.rows[0].cells[0].text = "Итого:"
            table.rows[0].cells[1].text = data["total"]
            table.rows[1].cells[0].text = f"НДС {data['vat']}%"
            table.rows[1].cells[1].text = data["total_vat"]
            table.rows[2].cells[0].text = "Всего с учетом НДС:"
            table.rows[2].cells[1].text = data["total_with_vat"]
            table.rows[3].cells[0].text = "Разница с заказом:"
            table.rows[3].cells[1].text = data["order_difference"]
            document.add_paragraph()
            for total_as_words in data["totals_as_words"]:
                document.add_paragraph(total_as_words)
            document.add_paragraph()
            for change in data["changes"]:
                document.add_paragraph(change)
            document.save(path)


class WordFormat2Converter(Converter):
    def convert(self, path: pathlib.Path, /, data: ConvertData) -> None:
        try:
            data["current_date"] = get_current_date()
            data["changes_in_smeta"] = False
            variants = ("сметному", "19.1", "19.2", "19.3", "20.1")
            if data["changes"]:
                for change in data["changes"]:
                    for variant in variants:
                        if variant in f'{change}':
                            data["changes_in_smeta"] = True

            import os

            from django.conf import settings
            from django.templatetags.static import static
            from docxtpl import DocxTemplate

            start_file_path = "/".join(
                os.path.abspath(__file__).replace("\\", "/").split("/")[:-2]
            ) + static("docx_templates/format2_start.docx")

            end_file_path = "/".join(
                os.path.abspath(__file__).replace("\\", "/").split("/")[:-2]
            ) + static("docx_templates/format2_end.docx")

            doc = DocxTemplate(start_file_path)
            data["br"] = "\n"
            # print(data)
            doc.render(data)
            doc.save(f"{str(path)[:-5]}start.docx")

            doc = DocxTemplate(end_file_path)
            doc.render(data)
            doc.save(f"{str(path)[:-5]}end.docx")

            combine_docx(
                f"{str(path)[:-5]}start.docx",
                f"{str(path)[:-5]}end.docx",
                data["smeta"],
                path,
            )

        except:
            traceback.print_exc()
            document: docx.document.Document = Document()
            rows = data["rows"]
            table = document.add_table(rows=len(rows), cols=7)
            table.style = "Table Grid"
            for i, row in enumerate(rows):
                for j, column in enumerate(row):
                    table.rows[i].cells[j].text = str(column)
            document.add_paragraph()
            table = document.add_table(rows=4, cols=2)
            table.style = "Table Grid"
            table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
            table.rows[0].cells[0].text = "Итого:"
            table.rows[0].cells[1].text = data["total"]
            table.rows[1].cells[0].text = f"НДС {data['vat']}%"
            table.rows[1].cells[1].text = data["total_vat"]
            table.rows[2].cells[0].text = "Всего с учетом НДС:"
            table.rows[2].cells[1].text = data["total_with_vat"]
            table.rows[3].cells[0].text = "Разница с заказом:"
            table.rows[3].cells[1].text = data["order_difference"]
            document.add_paragraph()
            for total_as_words in data["totals_as_words"]:
                document.add_paragraph(total_as_words)
            document.add_paragraph()
            for change in data["changes"]:
                document.add_paragraph(change)
            document.save(path)


class HTMLConverter(Converter):
    def convert(self, path: pathlib.Path, /, data: ConvertData) -> None:
        with open(path, "w", encoding="utf-8") as file:
            template: Template = get_template("p1/tester_atp_avr/preview.html")

            table = bs4.Tag(name="table", attrs={"class": "table table-bordered"})
            thead = bs4.Tag(name="thead")
            tr = bs4.Tag(name="tr")

            for column in data["rows"][0]:
                tag = bs4.Tag(name="th", attrs={"scope": "col", "class": "text-center"})
                tag.append(column)
                tr.append(tag)

            thead.append(tr)
            table.append(thead)
            tbody = bs4.Tag(name="tbody")
            for row in data["rows"][1:]:
                tr = bs4.Tag(
                    name="tr", attrs={"class": data["avr_rows"][row[1]].classes}
                )
                for index, column in enumerate(row):
                    tag = (
                        bs4.Tag(
                            name="th",
                            attrs={"scope": "row"},
                        )
                        if index == 0
                        else bs4.Tag(name="td")
                    )
                    tag.attrs["style"] = "height: 0;"
                    if index != 2:
                        content = bs4.Tag(
                            name="div",
                            attrs={
                                "style": "display: flex; align-items: center; justify-content: center; height: 100%; text-align: center;"
                            },
                        )
                        content.append(column)
                        tag.append(content)
                    else:
                        tag.append(column)
                    tr.append(tag)
                tbody.append(tr)
            table.append(tbody)
            file.write(
                template.render(
                    {
                        "table": str(table),
                        "data": data,
                    }
                )
            )


def get_tables_from_pdf(bytes_io: io.BytesIO, line_scale: int) -> TableList | None:
    fd, path = tempfile.mkstemp(".pdf")
    text = ""
    with os.fdopen(fd, "wb") as temporary_file:
        temporary_file.write(bytes_io.read())
    try:
        table_list = camelot.read_pdf(
            path,
            pages="all",
            split_text=True,
            line_scale=line_scale,
            strip_text="\n",
        )

        import fitz

        doc = fitz.open(path)

        for page_num in range(doc.page_count):
            page = doc[page_num]
            text += page.get_text()

        doc.close()

    except IndexError:
        return None
    return table_list, text


# def get_text_from_pdf(bytes_io: io.BytesIO, line_scale: int, **kwargs) -> str:
#     fd, path = tempfile.mkstemp(".pdf")
#     with os.fdopen(fd, "wb") as temporary_file:
#         temporary_file.write(bytes_io.read())
#     try:
#         text = camelot.read_pdf(
#             path,
#             pages="all",
#             split_text=True,
#             line_scale=line_scale,
#             strip_text="\n",
#         )
#     except IndexError:
#         return None
#     # os.remove(path)
#     return text


def search_additional_words(text):
    rejexes = {
        "bs_order": r"\S*\/[^\s\/]+\S*\/\S*",
        "bs_date": r"\b(?:\d{1,2}\s(?:январ[ья]|феврал[ья]|марта?|апрел[ья]|ма[йя]|июн[ья]|июл[ья]|август[а]|сентябр[ья]|октябр[ья]|ноябр[ья]|декабр[ья])\s\d{4}г?\.?|\d{2}\.\d{2}\.\d{2}(?:\d{2})?г?\.?)\b",
        "bs_number": r'\b(?:БС|BC|BS)[\s"\'"]*[\w\d_]+[\s"\'"]*\b',
        "bs_name": r"\b(?:[A-Z]+_){1,}[A-Z][a-z]*\b",
        "bs1_name": r"\b(?:[A-Z]+_){1,}[A-Z][a-z]*\b - \b(?:[A-Z]+_){1,}[A-Z][a-z]*\b",
        "work_type": r"\b[a-zA-Zа-яА-Я]+ работ(?:ы|а|)\b",
        "bs_address": r"адресу:\s*(.*?)В",
    }

    results = {}
    results["contract_number_and_date"] = None
    for rejex in rejexes:
        matches = re.findall(rejexes[rejex], text)
        if matches:

            results[rejex] = matches[0]
            if rejex == "bs_date" and len(matches) > 1:
                results["contract_number_and_date"] = find_closest_object(matches[1], ContratcNumberAndDate)
            elif rejex == "work_type" and len(matches) > 1:
                results["work_type"] = find_closest_object(matches[0], WorkType)
        else:
            results[rejex] = None
        if " - " in f"{results[rejex]}":
            results[rejex] = results[rejex].split(" - ")[1]

    # bs_name
    # bs1_name
    results["bs_address"] = ""
    return results


class TableDF1:
    df: pd.DataFrame

    def __init__(self, table: list[pd.DataFrame]):
        self.df = table


def get_tables_from_docx(bytes_io: io.BytesIO, line_scale: int):
    doc = Document(bytes_io)
    all_text = ""
    all_tables = []

    for paragraph in doc.paragraphs:
        all_text += paragraph.text + "\n"

    for table in doc.tables:
        current_table = pd.DataFrame(
            columns=[cell.text.strip() for cell in table.rows[0].cells]
        )

        for row in table.rows[1:]:
            current_row = [cell.text.strip() for cell in row.cells]
            current_table = pd.concat(
                [
                    current_table,
                    pd.DataFrame([current_row], columns=current_table.columns),
                ],
                ignore_index=True,
            )

        all_tables.append(TableDF1(current_table))

    return all_tables, all_text


import pandas as pd
import io
from openpyxl import load_workbook
from typing import Any, List, Tuple, Dict

def extract_data(file_path):
    wb = load_workbook(file_path, data_only=True)
    sheet = wb.active
    start_row = None
    end_row = None
    #
    # Find the start row (where column B is 1 and column C is 2)
    for row in sheet.iter_rows(min_row=1, max_row=sheet.max_row, min_col=2, max_col=3):
        b_value = row[0].value  # Column B (index 0)
        c_value = row[1].value  # Column C (index 1)
        if b_value == 1 and c_value == 2:
            start_row = row[0].row
            break
    #
    # Find the end row (where column B is "Всего по смете:")
    for row in sheet.iter_rows(min_row=1, max_row=sheet.max_row, min_col=3, max_col=3):
        if row[0].value and "Всего по смете:" in str(row[0].value):
            end_row = row[0].row
            break
    #
    if start_row is None:
        print("Could not find the required start or rows.")
        return []
    if end_row is None:
        print("Could not find the required end rows.")
        return []
    # Extract data from columns B to G between start_row and end_row (inclusive)
    data = []
    for row in sheet.iter_rows(min_row=start_row + 1, max_row=end_row, min_col=2, max_col=7):
        row_data = [cell.value for cell in row]
        data.append(row_data)
    return data

def get_numbered_data(data):
    numbered_data = []
    current_table = []
    current_number = 1
    #
    for row in data:
        first_cell = row[0]
        if isinstance(first_cell, int) and first_cell == current_number:
            current_table.append(row)
            current_number += 1
        else:
            if current_table:
                numbered_data.append(current_table)
                current_table = []
            current_number = 1
            if isinstance(first_cell, int) and first_cell == current_number:
                current_table.append(row)
                current_number += 1
    #
    if current_table:
        numbered_data.append(current_table)
    #
    return numbered_data


def extract_additional_data(data):
    search_words = ['итого материалы', 'итого по смете', 'ндс 12%', 'всего по смете']
    result = {}
    for i in data:
        if i[1]:  # Check if i[1] is not None
            for w in search_words:
                if w in i[1].lower():
                    result[w] = i[-1]
    return result

def extract_totals(data):
    search_words = {
        'итого работа': None,
        'итого материалы': None,
        'ндс 12%': None,
        'всего по смете': None,
        'итого по смете': None,
    }

    for row in data:
        if row[1]:  # Check if row[1] is not None
            for key in search_words.keys():
                if key in row[1].lower():
                    search_words[key] = row[-1]

    return (search_words['итого работа'],
            search_words['итого материалы'],
            search_words['ндс 12%'],
            search_words['итого по смете'],
            search_words['всего по смете'])


def parse_html_table(html):
    soup = bs4.BeautifulSoup(html, 'html.parser')
    print(soup.text)
    table = soup.find('table')
    extracted_data = []
    for row in table.find_all('tr'):
        cells = row.find_all('td')
        row_data = []
        for cell in cells:
            text = cell.get_text(strip=True)
            try:
                row_data.append(int(text))
            except ValueError:
                row_data.append(text)
        extracted_data.append(row_data)
    numbered_data = get_numbered_data(extracted_data)
    numbered_data_0 = numbered_data[0] if numbered_data else []
    numbered_data_1 = numbered_data[1] if len(numbered_data) > 1 else []
    total_work, total_materials, total_vat, total_overall, total_with_vat = extract_totals(extracted_data)
    return numbered_data_0, numbered_data_1, total_work, total_materials, total_vat, total_overall, total_with_vat

def parse_html_table_test(html):
    soup = bs4.BeautifulSoup(html, 'html.parser')
    print(soup.text)
    table = soup.find('table')
    extracted_data = []
    if table:
        for row in table.find_all('tr'):
            cells = row.find_all('td')
            row_data = []
            for cell in cells:
                text = cell.get_text(strip=True)
                try:
                    row_data.append(int(text))
                except ValueError:
                    row_data.append(text)
            extracted_data.append(row_data)
        return extracted_data
    else:
        return []
def calculate_additional_data(avr_work, avr_materials):
    total_work=0
    total_materials=0
    total_overall=0
    total_vat=0
    total_with_vat=0
    total_work=sum([ float(i[-1].replace(' ', '').replace(',', '.')) for i in avr_work ])
    total_materials=sum([ float(i[-1].replace(' ', '').replace(',', '.')) for i in avr_materials ])
    total_overall=total_work + total_materials
    total_vat = 0.12 * total_overall
    total_with_vat=total_overall+total_vat
    return total_work,total_materials,total_overall,total_vat, total_with_vat



def extract_data_between_indices(df: pd.DataFrame, start_index: int, end_index: int) -> List[List[Any]]:
    return df.iloc[start_index:end_index, 1:7].values.tolist()

def extract_data_between_indices_with_numbers(df: pd.DataFrame, start_index: int, end_index: int) -> List[List[Any]]:
    data = df.iloc[start_index:end_index, 1:7].values.tolist()
    pprint(df)
    filtered_data = [row for row in data if isinstance(row[0], (int, float))]
    return filtered_data if filtered_data else []

def find_indices(file_path):
    df = pd.read_excel(file_path)
    indices_value_2 = df.index[df.iloc[:, 2] == 2].tolist()
    indices_materials = df.index[df.iloc[:, 2].astype(str).str.contains("Материал", case=True, na=False)].tolist()
    indices_total_materials = df.index[
        df.iloc[:, 2].astype(str).str.contains("Итого материалы", case=False, na=False)].tolist()
    return {
        "indices_value_2": indices_value_2,
        "indices_materials": indices_materials,
        "indices_total_materials": indices_total_materials
    }

def get_tables_as_html_from_excel(file: io.BytesIO) -> Tuple[List[str], List[List[Any]], Dict[str, Any]]:
    try:
        wb = load_workbook(file, data_only=True)
        sheet = wb.active

        # Convert the sheet to a DataFrame
        df = pd.DataFrame(sheet.values)

        # Find indices
        indices = find_indices(file)
        print(indices)
        indices_value_2 = indices["indices_value_2"]
        indices_materials = indices["indices_materials"]
        indices_total_materials = indices["indices_total_materials"]

        if not indices_value_2 or not indices_materials or not indices_total_materials:
            raise ValueError("Required indices not found in the Excel file.")

        # Assuming the first index in the lists are the relevant ones for defining the ranges
        start_index = indices_value_2[0]
        middle_index = indices_materials[0]
        end_index = indices_total_materials[0]

        # Extract data for the first table
        first_table_data = extract_data_between_indices_with_numbers(df, start_index, middle_index)

        # Extract data for the second table, keeping only rows where the second column is a number
        second_table_data = extract_data_between_indices_with_numbers(df, middle_index, end_index)

        # Create numbered data (optional: modify according to your logic)
        numbered_data = [first_table_data, second_table_data]

        # Convert tables to HTML
        tables = []
        for table in numbered_data:
            df_table = pd.DataFrame(table)
            tables.append(
                df_table.to_html(
                    header=False, index=False, classes="table table-bordered", border=1
                )
            )

        return tables, numbered_data, {}
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise


def get_tables_as_html_from_pdf(bytes_io: io.BytesIO, **kwargs) -> list[str] | None:
    table_list, text = get_tables_from_pdf(bytes_io, **kwargs)

    if table_list is None:
        return None
    tables = []
    for table in table_list:
        tables.append(
            table.df.to_html(
                header=False, index=False, classes="table table-bordered", border=1
            )
        )
    return tables, text


def get_tables_as_html_from_docx(bytes_io: io.BytesIO, **kwargs) -> list[str] | None:
    table_list, text = get_tables_from_docx(bytes_io, **kwargs)
    if table_list is None:
        return None
    tables = []
    for table in table_list:
        tables.append(
            table.df.to_html(
                header=False, index=False, classes="table table-bordered", border=1
            )
        )
    return tables, text


def get_values_from_pdf_text(text) -> list[str] | None:
    if text is None:
        return {}
    print(text)
    additional_values = search_additional_words(text)
    return additional_values


@dataclass(slots=True)
class ATPRow:
    tcp: str
    name: str
    measuring_unit: str
    count: float | str
    due_date: datetime.datetime | str
    # услуги/тмц (опционально)
    service_or_tmc: str | None


def validate_atp(html: str | None = None) -> tuple[dict[str, ATPRow], list[str]]:
    errors = []
    atp_rows = {}

    if not html:
        errors.append(_("Вставьте таблицу"))
    else:
        bs = bs4.BeautifulSoup(html, "html.parser")
        print("start")
        print(bs.text)
        print("end")
        tables = bs.select("table")
        if tables:
            table = tables[0]
            for row in table.select("tr"):
                columns = row.select("td")
                datetime_ = columns[-1].text.strip()
                try:
                    datetime_ = datetime.datetime.strptime(datetime_, "%d.%m.%Y")
                except ValueError:
                    pass
                try:
                    if "услуга" == columns[3].text.lower().strip() or "ТМЦ" == columns[3].text.lower().strip():
                        atp_rows[columns[1].text.strip()] = ATPRow(
                            tcp=columns[1].text.strip(),
                            name=columns[2].text.strip(),
                            service_or_tmc=columns[3].text.strip(),
                            measuring_unit=columns[4].text.strip(),
                            count=float(columns[5].text.strip().replace(",", ".")),
                            due_date=datetime_,
                        )
                    else:
                        atp_rows[columns[1].text.strip()] = ATPRow(
                            tcp=columns[1].text.strip(),
                            name=columns[2].text.strip(),
                            measuring_unit=columns[3].text.strip(),
                            count=float(columns[4].text.strip().replace(",", ".")),
                            due_date=datetime_,
                            service_or_tmc=None,
                        )
                except (IndexError, ValueError):
                    errors.append(_("Данные не верные"))
                    break
        else:
            errors.append(_("Таблица не найдена"))
    return atp_rows, errors


@dataclass(slots=True)
class AVRRow(ATPRow):
    price: float | str
    added: bool
    classes: str = ""


def validate_avr(html: str | None = None) -> tuple[dict[str, AVRRow], list[str]]:
    errors = []
    avr_rows = {}
    if not html:
        errors.append(_("Вставьте таблицу"))
    else:
        bs = bs4.BeautifulSoup(html, "html.parser")

        tables = bs.select("table")
        if tables:
            table = tables[0]
            for row in table.select("tr"):
                columns = row.select("td")
                datetime_ = columns[6].text.strip()
                try:
                    datetime_ = datetime.datetime.strptime(datetime_, "%d.%m.%Y")
                except ValueError:
                    pass
                try:
                    if "услуга" == columns[3].text.lower().strip() or "ТМЦ" == columns[3].text.lower().strip():
                        avr_rows[columns[1].text.strip()] = AVRRow(
                            tcp=columns[1].text.strip(),
                            name=columns[2].text.strip(),
                            service_or_tmc=columns[3].text.strip(),
                            measuring_unit=columns[4].text.strip(),
                            count=float(columns[5].text.strip().replace(",", ".")),
                            due_date=datetime_,
                            price=float(remove_spaces(columns[6].text).replace(",", ".")),
                            added="data-new-row" in row.attrs,

                        )
                    else:
                        avr_rows[columns[1].text.strip()] = AVRRow(
                            tcp=columns[1].text.strip(),
                            name=columns[2].text.strip(),
                            measuring_unit=columns[3].text.strip(),
                            count=float(columns[4].text.strip().replace(",", ".")),
                            due_date=datetime_,
                            price=float(remove_spaces(columns[5].text).replace(",", ".")),
                            added="data-new-row" in row.attrs,
                            service_or_tmc=None,
                        )
                except (IndexError, ValueError):
                    errors.append(_("Данные не верные"))
                    break
        else:
            errors.append(_("Таблица не найдена"))

    return avr_rows, errors


class TCPCategoryRow(SimpleNamespace):
    tcp_category_id: int
    name: str
    note: str
    tcp_rows: list["TCPRow"]


class TCPRow(SimpleNamespace):
    tcp_id: int
    name: str
    note: str
    measuring_unit: int
    price: int


class PriceColumn(models.TextChoices):
    BASE_PRICE = "column_1", _("Базовая цена ТЦП, тг за ед., без НДС")
    BASE_PRICE_AFTER_DISCOUNT = (
        "column_2",
        _("Стоимость работ после скидки, тг без НДС"),
    )


def parse_1(
        worksheet: openpyxl.worksheet.worksheet.Worksheet, price_column: PriceColumn
) -> list[TCPCategoryRow]:
    tcp_category_rows = []

    for row in worksheet.iter_rows(min_row=2):
        tcp = row[1].value

        if tcp is None:
            continue
        tcp = str(tcp).split(".")
        if len(tcp) == 1:
            tcp_category_rows.append(
                TCPCategoryRow(
                    tcp_category_id=int(tcp[0]),
                    name=row[2].value,
                    note=row[5].value,
                    tcp_rows=[],
                )
            )
        else:
            assert len(tcp_category_rows) != 0
            price = (
                (0 if row[6].value is None else int(row[6].value))
                if price_column == PriceColumn.BASE_PRICE
                else (0 if row[7].value is None else int(row[7].value))
            )
            tcp_category_rows[-1].tcp_rows.append(
                TCPRow(
                    tcp_id=int(tcp[1]),
                    name=row[2].value,
                    note=row[5].value,
                    measuring_unit=row[4].value,
                    price=price,
                )
            )
    return tcp_category_rows


def parse_2(
        worksheet: openpyxl.worksheet.worksheet.Worksheet, _: PriceColumn
) -> list[TCPCategoryRow]:
    tcp_category_rows = []
    for row in worksheet.iter_rows(min_row=6):
        if row[0].value is None:
            if match := re.match(r"(\d+)\.(.+)", row[1].value):
                tcp_category_rows.append(
                    TCPCategoryRow(
                        tcp_category_id=int(match[1]),
                        name=match[2],
                        note="",
                        tcp_rows=[],
                    )
                )
        else:
            assert len(tcp_category_rows) != 0
            tcp_category_id, tcp_id = map(int, row[0].value.split("."))
            tcp_category = tcp_category_rows[-1]
            if tcp_category.tcp_category_id != tcp_category_id:
                tcp_category.tcp_category_id = tcp_category_id
            price = int(row[3].value)
            tcp_category.tcp_rows.append(
                TCPRow(
                    tcp_id=tcp_id,
                    name=row[1].value,
                    measuring_unit=row[1].value,
                    price=price,
                    note=row[4].value,
                )
            )
    return tcp_category_rows


def parse_3(
        worksheet: openpyxl.worksheet.worksheet.Worksheet, price_column: PriceColumn
):
    worksheet.delete_rows(1, 2)
    return parse_1(worksheet, price_column)


parsers: list[
    collections.abc.Callable[
        [openpyxl.worksheet.worksheet.Worksheet, PriceColumn], list[TCPCategoryRow]
    ]
] = [parse_1, parse_2, parse_3]


class LoadTCPError(Exception):
    pass


class ApplyError(LoadTCPError):
    pass


def apply_parsers(
        worksheet: openpyxl.worksheet.worksheet.Worksheet, price_column: PriceColumn
) -> list[TCPCategoryRow]:
    for index, parse in enumerate(parsers):
        logging.info(f"{index}. Applying the {parse.__qualname__} parser.")
        try:
            tcp_category_rows = parse(worksheet, price_column)
        except Exception as exception:
            logging.info(f"Error in applying: {exception}")
        else:
            return tcp_category_rows
    raise ApplyError


def load_tcp(tcp_file: TCPFile, price_column: PriceColumn) -> None:
    workbook = load_workbook(tcp_file.file.file, data_only=True)
    worksheet = workbook.active
    tcp_category_rows: list[TCPCategoryRow] = apply_parsers(worksheet, price_column)
    for tcp_category_row in tcp_category_rows:
        logging.info(
            f"Creating TCPCategory with: {tcp_category_row.tcp_category_id} {tcp_category_row.name}, note={tcp_category_row.note}."
        )
        tcp_category = TCPCategory.objects.create(
            tcp_category_id=tcp_category_row.tcp_category_id,
            name=tcp_category_row.name,
            note=tcp_category_row.note,
            tcp_file=tcp_file,
        )
        for tcp_row in tcp_category_row.tcp_rows:
            logging.info(
                f"Creating TCP with: {tcp_category_row.tcp_category_id}.{tcp_row.tcp_id} {tcp_row.name}, note={tcp_row.note}, price={tcp_row.price}."
            )
            TCP.objects.create(
                tcp_id=tcp_row.tcp_id,
                tcp_category=tcp_category,
                name=tcp_row.name,
                note=tcp_row.note,
                price=tcp_row.price,
            )


def get_tables_as_list(html_content):
    if html_content is None:
        return None

    soup = bs4.BeautifulSoup(html_content, "html.parser")

    tables_as_list = []
    for row in soup.find_all("tr"):
        tables_as_list.append([cell.get_text() for cell in row.find_all("td")])

    if len(tables_as_list) == 0:
        print("No tables found")
        return None
    return tables_as_list


def get_tables_as_list_smeta(bytes_io: io.BytesIO):
    if bytes_io:
        return bytes_io
    return None


def combine_docx(template_start, template_end, smeta_BytesIO, output_file_path):
    doc1 = Document(template_start)
    doc2 = Document(template_end)
    doc3 = Document(smeta_BytesIO)

    docX = doc1

    if smeta_BytesIO:
        for element in doc3.element.body:
            docX.element.body.append(element)

    for element in doc2.element.body:
        docX.element.body.append(element)

    combined_doc = Document()
    combined_doc = docX
    combined_doc.save(output_file_path)


def find_closest_object(input_text, object_class):
    all_objects = object_class.objects.all()

    if not all_objects:
        return None

    closest_object = min(
        all_objects, key=lambda x: distance(input_text, x.__str__())
    )

    return closest_object.name


def get_current_date():
    # Получаем текущую дату
    current_date = datetime.datetime.now()

    # Месяцы на русском
    russian_months = [
        'января', 'февраля', 'марта', 'апреля',
        'мая', 'июня', 'июля', 'августа',
        'сентября', 'октября', 'ноября', 'декабря'
    ]

    # Форматируем дату
    formatted_date = current_date.strftime(f'«%d» {russian_months[current_date.month - 1]} %Y г.')

    return formatted_date
